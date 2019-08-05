import docx

class Word_Text_Extraction():
    def __init__(self, file_name):
        self.file_name = file_name
    
    def text_extraction(self):
        document = docx.Document(self.file_name)
        all_text = []

        for i in range(len(document.paragraphs)):
            paragraph_text = document.paragraphs[i].text
            all_text.append(paragraph_text)
        
        word_document = " ".join(all_text)
        print(word_document)
    
    def execute(self):
        self.text_extraction()

parse = Word_Text_Extraction("McMaster_WebDeveloper_Task.docx")
parse.execute()